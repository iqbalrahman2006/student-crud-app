"""
PART 3: Add comprehensive workflows, DDL with triggers, and constraint documentation
"""

from docx import Document
from docx.shared import Pt
import os

doc = Document('STUDENTDB_INFORMATIONAL_DOCUMENTATION.docx')

def add_section_header(doc, title):
    return doc.add_heading(title, level=1)

def add_subsection(doc, title):
    return doc.add_heading(title, level=2)

def add_code_block(doc, code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return p

# ==================== SECTION 5: COMPLETE DDL ====================
add_section_header(doc, '5. COMPLETE DDL WITH CONSTRAINTS')

add_subsection(doc, '5.1 Students Table - Complete DDL')

students_ddl = """CREATE TABLE Students (
    student_id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
    name VARCHAR(255) NOT NULL CHECK (LENGTH(TRIM(name)) >= 2 AND LENGTH(TRIM(name)) <= 255),
    email VARCHAR(255) NOT NULL UNIQUE CHECK (email ~* '^[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\\.[A-Z|a-z]{2,}$'),
    phone VARCHAR(20) CHECK (phone IS NULL OR phone ~* '^[\\d\\s\\-+()]{10,20}$'),
    course VARCHAR(100) CHECK (course IS NULL OR LENGTH(TRIM(course)) >= 2),
    department VARCHAR(100),
    status VARCHAR(20) NOT NULL DEFAULT 'Active' 
        CHECK (status IN ('Active', 'Inactive', 'Suspended', 'Graduated', 'Expelled')),
    enrollmentDate TIMESTAMP NOT NULL DEFAULT NOW(),
    gpa DECIMAL(4,2) CHECK (gpa IS NULL OR (gpa >= 0.00 AND gpa <= 10.00)),
    city VARCHAR(100),
    country VARCHAR(100) DEFAULT 'Pakistan',
    zipCode VARCHAR(20),
    address TEXT,
    guardianName VARCHAR(255),
    guardianPhone VARCHAR(20),
    emergencyContact VARCHAR(20),
    studentCategory VARCHAR(50) 
        CHECK (studentCategory IS NULL OR studentCategory IN ('Regular', 'Exchange', 'International', 'Part-time')),
    scholarshipStatus VARCHAR(50) 
        CHECK (scholarshipStatus IS NULL OR scholarshipStatus IN ('None', 'Partial', 'Full', 'Merit', 'Need-based')),
    bloodGroup VARCHAR(3) 
        CHECK (bloodGroup IS NULL OR bloodGroup IN ('A+', 'A-', 'B+', 'B-', 'AB+', 'AB-', 'O+', 'O-')),
    hostelRequired BOOLEAN DEFAULT FALSE,
    transportMode VARCHAR(50),
    libraryCardNumber VARCHAR(50) UNIQUE,
    maxBooksAllowed INTEGER DEFAULT 5 CHECK (maxBooksAllowed >= 0 AND maxBooksAllowed <= 20),
    currentBooksIssued INTEGER DEFAULT 0 CHECK (currentBooksIssued >= 0),
    totalBooksIssuedLifetime INTEGER DEFAULT 0 CHECK (totalBooksIssuedLifetime >= 0),
    totalFinesAccumulated DECIMAL(10,2) DEFAULT 0.00 CHECK (totalFinesAccumulated >= 0),
    accountStatus VARCHAR(20) DEFAULT 'Good Standing',
    lastBorrowDate TIMESTAMP,
    lastReturnDate TIMESTAMP,
    notes TEXT,
    createdAt TIMESTAMP NOT NULL DEFAULT NOW(),
    updatedAt TIMESTAMP NOT NULL DEFAULT NOW(),
    isDeleted BOOLEAN DEFAULT FALSE
);

-- Indexes
CREATE UNIQUE INDEX idx_students_email ON Students(email) WHERE isDeleted = FALSE;
CREATE UNIQUE INDEX idx_students_library_card ON Students(libraryCardNumber) 
    WHERE libraryCardNumber IS NOT NULL AND isDeleted = FALSE;
CREATE INDEX idx_students_status ON Students(status) WHERE isDeleted = FALSE;
CREATE INDEX idx_students_department ON Students(department) WHERE isDeleted = FALSE;
CREATE INDEX idx_students_enrollment_date ON Students(enrollmentDate DESC);
CREATE INDEX idx_students_gpa ON Students(gpa DESC) WHERE gpa IS NOT NULL;
CREATE INDEX idx_students_account_status ON Students(accountStatus) WHERE isDeleted = FALSE;

-- Trigger: Prevent modification of immutable fields
CREATE OR REPLACE FUNCTION prevent_student_immutable_fields()
RETURNS TRIGGER AS $$
BEGIN
    IF OLD.email IS DISTINCT FROM NEW.email THEN
        RAISE EXCEPTION 'Student email is immutable and cannot be changed';
    END IF;
    IF OLD.enrollmentDate IS DISTINCT FROM NEW.enrollmentDate THEN
        RAISE EXCEPTION 'Enrollment date is immutable and cannot be changed';
    END IF;
    IF OLD.libraryCardNumber IS DISTINCT FROM NEW.libraryCardNumber 
       AND OLD.libraryCardNumber IS NOT NULL THEN
        RAISE EXCEPTION 'Library card number is immutable once assigned';
    END IF;
    NEW.updatedAt = NOW();
    RETURN NEW;
END;
$$ LANGUAGE plpgsql;

CREATE TRIGGER trg_student_immutability
BEFORE UPDATE ON Students
FOR EACH ROW
EXECUTE FUNCTION prevent_student_immutable_fields();

-- Trigger: Validate books count
CREATE OR REPLACE FUNCTION validate_student_books_count()
RETURNS TRIGGER AS $$
BEGIN
    IF NEW.currentBooksIssued > NEW.maxBooksAllowed THEN
        RAISE EXCEPTION 'Current books issued (%) exceeds maximum allowed (%)', 
            NEW.currentBooksIssued, NEW.maxBooksAllowed;
    END IF;
    IF NEW.currentBooksIssued < 0 THEN
        RAISE EXCEPTION 'Current books issued cannot be negative';
    END IF;
    RETURN NEW;
END;
$$ LANGUAGE plpgsql;

CREATE TRIGGER trg_student_books_validation
BEFORE INSERT OR UPDATE ON Students
FOR EACH ROW
EXECUTE FUNCTION validate_student_books_count();
"""

add_code_block(doc, students_ddl)

doc.add_page_break()

add_subsection(doc, '5.2 Books Table - Complete DDL')

books_ddl = """CREATE TABLE Books (
    book_id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
    title VARCHAR(500) NOT NULL CHECK (LENGTH(TRIM(title)) >= 2 AND LENGTH(TRIM(title)) <= 500),
    author VARCHAR(255) NOT NULL CHECK (LENGTH(TRIM(author)) >= 2 AND LENGTH(TRIM(author)) <= 255),
    isbn VARCHAR(50) NOT NULL UNIQUE CHECK (LENGTH(TRIM(isbn)) >= 10),
    publisher VARCHAR(255),
    publicationYear INTEGER CHECK (publicationYear IS NULL OR 
        (publicationYear >= 1800 AND publicationYear <= EXTRACT(YEAR FROM NOW()) + 1)),
    edition VARCHAR(50),
    genre VARCHAR(100),
    language VARCHAR(50) DEFAULT 'English',
    department VARCHAR(100) NOT NULL CHECK (department IN (
        'Computer Science', 'Electrical Engineering', 'Mechanical Engineering', 
        'Civil Engineering', 'Business Administration', 'Medicine', 'Literature', 
        'Philosophy', 'Mathematics', 'Physics', 'Chemistry', 'Biology', 'General'
    )),
    totalCopies INTEGER NOT NULL DEFAULT 1 CHECK (totalCopies >= 0 AND totalCopies <= 1000),
    checkedOutCount INTEGER NOT NULL DEFAULT 0 CHECK (checkedOutCount >= 0),
    availableCopies INTEGER GENERATED ALWAYS AS (totalCopies - checkedOutCount) STORED,
    reservedCount INTEGER DEFAULT 0 CHECK (reservedCount >= 0),
    damagedCount INTEGER DEFAULT 0 CHECK (damagedCount >= 0),
    lostCount INTEGER DEFAULT 0 CHECK (lostCount >= 0),
    status VARCHAR(20) NOT NULL DEFAULT 'Available' 
        CHECK (status IN ('Available', 'Out of Stock', 'Discontinued', 'Under Review', 'Damaged')),
    shelfLocation VARCHAR(50),
    callNumber VARCHAR(50),
    price DECIMAL(10,2) CHECK (price IS NULL OR price >= 0),
    purchaseDate TIMESTAMP,
    condition VARCHAR(20) DEFAULT 'New' CHECK (condition IN ('New', 'Good', 'Fair', 'Poor', 'Damaged')),
    pages INTEGER CHECK (pages IS NULL OR pages > 0),
    totalBorrowCount INTEGER DEFAULT 0 CHECK (totalBorrowCount >= 0),
    averageRating DECIMAL(3,2) CHECK (averageRating IS NULL OR (averageRating >= 0 AND averageRating <= 5.00)),
    lastBorrowedDate TIMESTAMP,
    lastReturnedDate TIMESTAMP,
    addedDate TIMESTAMP NOT NULL DEFAULT NOW(),
    createdAt TIMESTAMP NOT NULL DEFAULT NOW(),
    updatedAt TIMESTAMP NOT NULL DEFAULT NOW(),
    isDeleted BOOLEAN DEFAULT FALSE
);

-- Indexes
CREATE UNIQUE INDEX idx_books_isbn ON Books(isbn) WHERE isDeleted = FALSE;
CREATE INDEX idx_books_title ON Books USING gin(to_tsvector('english', title));
CREATE INDEX idx_books_author ON Books USING gin(to_tsvector('english', author));
CREATE INDEX idx_books_department ON Books(department) WHERE isDeleted = FALSE;
CREATE INDEX idx_books_status ON Books(status) WHERE isDeleted = FALSE;
CREATE INDEX idx_books_available ON Books(availableCopies) WHERE availableCopies > 0 AND isDeleted = FALSE;
CREATE INDEX idx_books_shelf_location ON Books(shelfLocation) WHERE shelfLocation IS NOT NULL;

-- Trigger: Prevent modification of immutable fields
CREATE OR REPLACE FUNCTION prevent_book_immutable_fields()
RETURNS TRIGGER AS $$
BEGIN
    IF OLD.isbn IS DISTINCT FROM NEW.isbn THEN
        RAISE EXCEPTION 'Book ISBN is immutable and cannot be changed';
    END IF;
    IF OLD.addedDate IS DISTINCT FROM NEW.addedDate THEN
        RAISE EXCEPTION 'Book added date is immutable and cannot be changed';
    END IF;
    NEW.updatedAt = NOW();
    RETURN NEW;
END;
$$ LANGUAGE plpgsql;

CREATE TRIGGER trg_book_immutability
BEFORE UPDATE ON Books
FOR EACH ROW
EXECUTE FUNCTION prevent_book_immutable_fields();

-- Trigger: Validate book counts
CREATE OR REPLACE FUNCTION validate_book_counts()
RETURNS TRIGGER AS $$
BEGIN
    IF NEW.checkedOutCount > NEW.totalCopies THEN
        RAISE EXCEPTION 'Checked out count (%) cannot exceed total copies (%)', 
            NEW.checkedOutCount, NEW.totalCopies;
    END IF;
    IF (NEW.checkedOutCount + NEW.damagedCount + NEW.lostCount) > NEW.totalCopies THEN
        RAISE EXCEPTION 'Sum of checked out, damaged, and lost cannot exceed total copies';
    END IF;
    RETURN NEW;
END;
$$ LANGUAGE plpgsql;

CREATE TRIGGER trg_book_counts_validation
BEFORE INSERT OR UPDATE ON Books
FOR EACH ROW
EXECUTE FUNCTION validate_book_counts();

-- Trigger: Auto-update status based on availability
CREATE OR REPLACE FUNCTION update_book_status()
RETURNS TRIGGER AS $$
BEGIN
    IF NEW.availableCopies = 0 AND NEW.status = 'Available' THEN
        NEW.status = 'Out of Stock';
    ELSIF NEW.availableCopies > 0 AND NEW.status = 'Out of Stock' THEN
        NEW.status = 'Available';
    END IF;
    RETURN NEW;
END;
$$ LANGUAGE plpgsql;

CREATE TRIGGER trg_book_status_update
BEFORE INSERT OR UPDATE ON Books
FOR EACH ROW
EXECUTE FUNCTION update_book_status();
"""

add_code_block(doc, books_ddl)

doc.add_page_break()

add_subsection(doc, '5.3 BorrowTransactions Table - Complete DDL')

transactions_ddl = """CREATE TABLE BorrowTransactions (
    transaction_id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
    studentId UUID NOT NULL REFERENCES Students(student_id) ON DELETE RESTRICT ON UPDATE CASCADE,
    bookId UUID NOT NULL REFERENCES Books(book_id) ON DELETE RESTRICT ON UPDATE CASCADE,
    issuedAt TIMESTAMP NOT NULL DEFAULT NOW(),
    dueDate TIMESTAMP NOT NULL,
    returnedAt TIMESTAMP,
    status VARCHAR(20) NOT NULL DEFAULT 'BORROWED' 
        CHECK (status IN ('BORROWED', 'RETURNED', 'OVERDUE', 'LOST', 'DAMAGED', 'RENEWED')),
    renewalCount INTEGER DEFAULT 0 CHECK (renewalCount >= 0 AND renewalCount <= 5),
    maxRenewalsAllowed INTEGER DEFAULT 3 CHECK (maxRenewalsAllowed >= 0),
    fineAmount DECIMAL(10,2) DEFAULT 0.00 CHECK (fineAmount >= 0),
    finePaid BOOLEAN DEFAULT FALSE,
    finePaymentDate TIMESTAMP,
    daysOverdue INTEGER DEFAULT 0 CHECK (daysOverdue >= 0),
    issuedBy UUID REFERENCES Users(user_id) ON DELETE SET NULL ON UPDATE CASCADE,
    returnedTo UUID REFERENCES Users(user_id) ON DELETE SET NULL ON UPDATE CASCADE,
    bookConditionAtIssue VARCHAR(20) DEFAULT 'Good',
    bookConditionAtReturn VARCHAR(20),
    damageCharge DECIMAL(10,2) DEFAULT 0.00 CHECK (damageCharge >= 0),
    createdAt TIMESTAMP NOT NULL DEFAULT NOW(),
    updatedAt TIMESTAMP NOT NULL DEFAULT NOW(),
    isDeleted BOOLEAN DEFAULT FALSE
);

-- Indexes
CREATE INDEX idx_borrow_student ON BorrowTransactions(studentId) WHERE isDeleted = FALSE;
CREATE INDEX idx_borrow_book ON BorrowTransactions(bookId) WHERE isDeleted = FALSE;
CREATE INDEX idx_borrow_status ON BorrowTransactions(status) WHERE isDeleted = FALSE;
CREATE INDEX idx_borrow_student_status ON BorrowTransactions(studentId, status) WHERE isDeleted = FALSE;
CREATE INDEX idx_borrow_book_status ON BorrowTransactions(bookId, status) WHERE isDeleted = FALSE;
CREATE INDEX idx_borrow_due_date ON BorrowTransactions(dueDate) 
    WHERE status IN ('BORROWED', 'OVERDUE') AND isDeleted = FALSE;
CREATE INDEX idx_borrow_overdue ON BorrowTransactions(dueDate, status) 
    WHERE status = 'BORROWED' AND dueDate < NOW() AND isDeleted = FALSE;

-- Trigger: Validate foreign key references
CREATE OR REPLACE FUNCTION validate_borrow_transaction_refs()
RETURNS TRIGGER AS $$
DECLARE
    student_exists BOOLEAN;
    book_exists BOOLEAN;
    book_available INTEGER;
BEGIN
    SELECT EXISTS(SELECT 1 FROM Students WHERE student_id = NEW.studentId AND isDeleted = FALSE) 
    INTO student_exists;
    IF NOT student_exists THEN
        RAISE EXCEPTION 'Referenced student does not exist or is deleted';
    END IF;
    
    SELECT EXISTS(SELECT 1 FROM Books WHERE book_id = NEW.bookId AND isDeleted = FALSE) 
    INTO book_exists;
    IF NOT book_exists THEN
        RAISE EXCEPTION 'Referenced book does not exist or is deleted';
    END IF;
    
    IF TG_OP = 'INSERT' THEN
        SELECT availableCopies INTO book_available FROM Books WHERE book_id = NEW.bookId;
        IF book_available <= 0 THEN
            RAISE EXCEPTION 'Book is not available for borrowing';
        END IF;
    END IF;
    
    RETURN NEW;
END;
$$ LANGUAGE plpgsql;

CREATE TRIGGER trg_borrow_transaction_validation
BEFORE INSERT OR UPDATE ON BorrowTransactions
FOR EACH ROW
EXECUTE FUNCTION validate_borrow_transaction_refs();

-- Trigger: Calculate overdue fine
CREATE OR REPLACE FUNCTION calculate_overdue_fine()
RETURNS TRIGGER AS $$
BEGIN
    IF NEW.status = 'RETURNED' AND NEW.returnedAt > NEW.dueDate THEN
        NEW.daysOverdue = EXTRACT(DAY FROM (NEW.returnedAt - NEW.dueDate));
        NEW.fineAmount = NEW.daysOverdue * 5.00;  -- Rs. 5 per day
    ELSIF NEW.status = 'OVERDUE' THEN
        NEW.daysOverdue = EXTRACT(DAY FROM (NOW() - NEW.dueDate));
        NEW.fineAmount = NEW.daysOverdue * 5.00;
    END IF;
    RETURN NEW;
END;
$$ LANGUAGE plpgsql;

CREATE TRIGGER trg_calculate_fine
BEFORE INSERT OR UPDATE ON BorrowTransactions
FOR EACH ROW
EXECUTE FUNCTION calculate_overdue_fine();
"""

add_code_block(doc, transactions_ddl)

doc.add_page_break()

add_subsection(doc, '5.4 LibraryAuditLogs Table - Complete DDL (IMMUTABLE)')

audit_ddl = """CREATE TABLE LibraryAuditLogs (
    log_id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
    action VARCHAR(50) NOT NULL CHECK (action IN (
        'BORROW', 'RETURN', 'RENEW', 'RESERVE', 'CANCEL_RESERVATION',
        'ADD_BOOK', 'UPDATE_BOOK', 'DELETE_BOOK',
        'ADD_STUDENT', 'UPDATE_STUDENT', 'DELETE_STUDENT',
        'FINE_PAID', 'EMAIL_SENT', 'OVERDUE_DETECTED'
    )),
    entityType VARCHAR(50) CHECK (entityType IN ('Student', 'Book', 'Transaction', 'Reservation', 'Fine', 'User')),
    entityId UUID,
    studentId UUID REFERENCES Students(student_id) ON DELETE SET NULL ON UPDATE CASCADE,
    bookId UUID REFERENCES Books(book_id) ON DELETE SET NULL ON UPDATE CASCADE,
    adminId UUID REFERENCES Users(user_id) ON DELETE SET NULL ON UPDATE CASCADE,
    timestamp TIMESTAMP NOT NULL DEFAULT NOW(),
    metadata JSONB,
    ipAddress VARCHAR(45),
    userAgent TEXT,
    success BOOLEAN DEFAULT TRUE,
    errorMessage TEXT,
    createdAt TIMESTAMP NOT NULL DEFAULT NOW()
);

-- Indexes
CREATE INDEX idx_audit_action ON LibraryAuditLogs(action);
CREATE INDEX idx_audit_timestamp ON LibraryAuditLogs(timestamp DESC);
CREATE INDEX idx_audit_student ON LibraryAuditLogs(studentId) WHERE studentId IS NOT NULL;
CREATE INDEX idx_audit_book ON LibraryAuditLogs(bookId) WHERE bookId IS NOT NULL;
CREATE INDEX idx_audit_admin ON LibraryAuditLogs(adminId) WHERE adminId IS NOT NULL;
CREATE INDEX idx_audit_entity ON LibraryAuditLogs(entityType, entityId) WHERE entityType IS NOT NULL;
CREATE INDEX idx_audit_action_timestamp ON LibraryAuditLogs(action, timestamp DESC);

-- Trigger: ENFORCE IMMUTABILITY - Prevent ANY modifications
CREATE OR REPLACE FUNCTION prevent_audit_log_modification()
RETURNS TRIGGER AS $$
BEGIN
    RAISE EXCEPTION 'Audit logs are immutable and cannot be modified or deleted. Operation: %', TG_OP;
END;
$$ LANGUAGE plpgsql;

CREATE TRIGGER trg_audit_log_immutability_update
BEFORE UPDATE ON LibraryAuditLogs
FOR EACH ROW
EXECUTE FUNCTION prevent_audit_log_modification();

CREATE TRIGGER trg_audit_log_immutability_delete
BEFORE DELETE ON LibraryAuditLogs
FOR EACH ROW
EXECUTE FUNCTION prevent_audit_log_modification();

COMMENT ON TABLE LibraryAuditLogs IS 'Immutable append-only audit trail. NO UPDATES OR DELETES ALLOWED.';
"""

add_code_block(doc, audit_ddl)

doc.add_page_break()

# Save progress
doc.save('STUDENTDB_INFORMATIONAL_DOCUMENTATION.docx')

file_size = os.path.getsize('STUDENTDB_INFORMATIONAL_DOCUMENTATION.docx')
file_size_kb = file_size / 1024

print(f"\n{'='*60}")
print(f"[PROGRESS] Added complete DDL with triggers")
print(f"{'='*60}")
print(f"Current Size: {file_size_kb:.2f} KB")
print(f"Target: 100 KB")
print(f"Progress: {(file_size_kb/100)*100:.1f}%")
print(f"{'='*60}\n")

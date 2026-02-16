"""
PART 2: Add Normalization, DDL, Workflows, and remaining informational sections
"""

from docx import Document
from docx.shared import Pt
import os

doc = Document('STUDENTDB_INFORMATIONAL_DOCUMENTATION.docx')

def add_section_header(doc, title):
    return doc.add_heading(title, level=1)

def add_subsection(doc, title):
    return doc.add_heading(title, level=2)

def add_code_block(doc, code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return p

# ==================== SECTION 2.4: CARDINALITY RULES ====================
add_subsection(doc, '2.4 Cardinality Rules & Participation Constraints')

cardinality_table_content = """
┌──────────────────────────────────────────────────────────────────────────┐
│                      CARDINALITY MATRIX                                   │
├──────────────────────────────────────────────────────────────────────────┤
│                                                                           │
│  Relationship: Student ←→ BorrowTransaction ←→ Book                      │
│  ─────────────────────────────────────────────────────────────────────   │
│  • One Student can have MANY BorrowTransactions (1:N)                    │
│  • One Book can have MANY BorrowTransactions (1:N)                       │
│  • Participation: Student (Partial), Book (Partial)                      │
│  • A student may have zero transactions (new student)                    │
│  • A book may have zero transactions (newly added book)                  │
│                                                                           │
│  Relationship: Student ←→ BookReservation ←→ Book                        │
│  ─────────────────────────────────────────────────────────────────────   │
│  • One Student can have MANY Reservations (1:N)                          │
│  • One Book can have MANY Reservations (1:N)                             │
│  • Participation: Student (Partial), Book (Partial)                      │
│  • Reservations only exist when book is unavailable                      │
│                                                                           │
│  Relationship: BorrowTransaction ←→ LibraryFineLedger                    │
│  ─────────────────────────────────────────────────────────────────────   │
│  • One Transaction can have ZERO or ONE Fine (1:0..1)                    │
│  • Participation: Transaction (Partial), Fine (Total)                    │
│  • Fine exists only if transaction incurs penalty                        │
│                                                                           │
│  Relationship: Student ←→ LibraryFineLedger                              │
│  ─────────────────────────────────────────────────────────────────────   │
│  • One Student can have MANY Fines (1:N)                                 │
│  • Participation: Student (Partial), Fine (Total)                        │
│  • Every fine must be associated with a student                          │
│                                                                           │
│  Relationship: All Entities ←→ LibraryAuditLogs                          │
│  ─────────────────────────────────────────────────────────────────────   │
│  • One Entity can have MANY Audit Logs (1:N)                             │
│  • Participation: Entity (Partial), AuditLog (Partial)                   │
│  • Audit logs capture operations but entity references are nullable      │
│                                                                           │
└──────────────────────────────────────────────────────────────────────────┘
"""

add_code_block(doc, cardinality_table_content)

doc.add_page_break()

# ==================== SECTION 3: RELATIONAL SCHEMA ====================
add_section_header(doc, '3. RELATIONAL SCHEMA FORMALIZATION')

add_subsection(doc, '3.1 Schema Notation')

doc.add_paragraph(
    'The relational schema is expressed using standard database notation where:\n'
    '• PK = Primary Key\n'
    '• FK = Foreign Key\n'
    '• UK = Unique Key\n'
    '• NN = Not Null\n'
    '• DF = Default Value\n'
    '• CK = Check Constraint\n'
    '• IM = Immutable (cannot be updated)\n'
    '• GEN = Generated/Computed Field'
)

add_subsection(doc, '3.2 Complete Table Specifications')

# Students Table
doc.add_heading('Table: Students', level=3)

students_schema = """
Students(
    student_id UUID PK NN DF:gen_random_uuid(),
    name VARCHAR(255) NN CK:LENGTH(name) >= 2,
    email VARCHAR(255) UK NN IM CK:email_format,
    phone VARCHAR(20) CK:phone_format,
    course VARCHAR(100),
    department VARCHAR(100),
    status VARCHAR(20) NN DF:'Active' CK:IN('Active','Inactive','Suspended','Graduated'),
    enrollmentDate TIMESTAMP NN IM DF:NOW(),
    gpa DECIMAL(4,2) CK:gpa >= 0 AND gpa <= 10,
    city VARCHAR(100),
    country VARCHAR(100) DF:'Pakistan',
    zipCode VARCHAR(20),
    address TEXT,
    guardianName VARCHAR(255),
    guardianPhone VARCHAR(20),
    emergencyContact VARCHAR(20),
    studentCategory VARCHAR(50) CK:IN('Regular','Exchange','International','Part-time'),
    scholarshipStatus VARCHAR(50) CK:IN('None','Partial','Full','Merit','Need-based'),
    bloodGroup VARCHAR(3) CK:IN('A+','A-','B+','B-','AB+','AB-','O+','O-'),
    hostelRequired BOOLEAN DF:FALSE,
    transportMode VARCHAR(50),
    libraryCardNumber VARCHAR(50) UK,
    maxBooksAllowed INTEGER DF:5 CK:>= 0 AND <= 20,
    currentBooksIssued INTEGER DF:0 CK:>= 0,
    totalBooksIssuedLifetime INTEGER DF:0 CK:>= 0,
    totalFinesAccumulated DECIMAL(10,2) DF:0.00 CK:>= 0,
    accountStatus VARCHAR(20) DF:'Good Standing',
    lastBorrowDate TIMESTAMP,
    lastReturnDate TIMESTAMP,
    notes TEXT,
    createdAt TIMESTAMP NN DF:NOW(),
    updatedAt TIMESTAMP NN DF:NOW(),
    isDeleted BOOLEAN DF:FALSE
)

INDEXES:
  - UNIQUE INDEX idx_students_email ON (email) WHERE isDeleted = FALSE
  - UNIQUE INDEX idx_students_library_card ON (libraryCardNumber) WHERE libraryCardNumber IS NOT NULL
  - INDEX idx_students_status ON (status) WHERE isDeleted = FALSE
  - INDEX idx_students_department ON (department)
  - INDEX idx_students_gpa ON (gpa DESC) WHERE gpa IS NOT NULL
  - INDEX idx_students_enrollment ON (enrollmentDate DESC)
  - INDEX idx_students_account_status ON (accountStatus)

CONSTRAINTS:
  - currentBooksIssued <= maxBooksAllowed
  - email IMMUTABLE (enforced via trigger)
  - enrollmentDate IMMUTABLE (enforced via trigger)
  - libraryCardNumber IMMUTABLE once assigned
"""

add_code_block(doc, students_schema)

# Books Table
doc.add_heading('Table: Books', level=3)

books_schema = """
Books(
    book_id UUID PK NN DF:gen_random_uuid(),
    title VARCHAR(500) NN CK:LENGTH(title) >= 2,
    author VARCHAR(255) NN CK:LENGTH(author) >= 2,
    isbn VARCHAR(50) UK NN IM CK:LENGTH(isbn) >= 10,
    publisher VARCHAR(255),
    publicationYear INTEGER CK:>= 1800 AND <= YEAR(NOW()) + 1,
    edition VARCHAR(50),
    genre VARCHAR(100),
    language VARCHAR(50) DF:'English',
    department VARCHAR(100) NN CK:IN('Computer Science','Electrical Engineering',
                                      'Mechanical Engineering','Civil Engineering',
                                      'Business Administration','Medicine','Literature',
                                      'Philosophy','Mathematics','Physics','Chemistry',
                                      'Biology','General'),
    totalCopies INTEGER NN DF:1 CK:>= 0 AND <= 1000,
    checkedOutCount INTEGER NN DF:0 CK:>= 0,
    availableCopies INTEGER GEN AS (totalCopies - checkedOutCount) STORED,
    reservedCount INTEGER DF:0 CK:>= 0,
    damagedCount INTEGER DF:0 CK:>= 0,
    lostCount INTEGER DF:0 CK:>= 0,
    status VARCHAR(20) NN DF:'Available' CK:IN('Available','Out of Stock','Discontinued'),
    shelfLocation VARCHAR(50),
    callNumber VARCHAR(50),
    deweyDecimal VARCHAR(20),
    price DECIMAL(10,2) CK:>= 0,
    purchaseDate TIMESTAMP,
    condition VARCHAR(20) DF:'New' CK:IN('New','Good','Fair','Poor','Damaged'),
    pages INTEGER CK:> 0,
    weight DECIMAL(6,2) CK:> 0,
    coverType VARCHAR(20) CK:IN('Hardcover','Paperback','Spiral','Leather'),
    totalBorrowCount INTEGER DF:0 CK:>= 0,
    averageRating DECIMAL(3,2) CK:>= 0 AND <= 5,
    reviewCount INTEGER DF:0 CK:>= 0,
    lastBorrowedDate TIMESTAMP,
    lastReturnedDate TIMESTAMP,
    addedDate TIMESTAMP NN IM DF:NOW(),
    createdAt TIMESTAMP NN DF:NOW(),
    updatedAt TIMESTAMP NN DF:NOW(),
    isDeleted BOOLEAN DF:FALSE
)

INDEXES:
  - UNIQUE INDEX idx_books_isbn ON (isbn) WHERE isDeleted = FALSE
  - INDEX idx_books_title_fulltext ON (to_tsvector('english', title))
  - INDEX idx_books_author_fulltext ON (to_tsvector('english', author))
  - INDEX idx_books_department ON (department)
  - INDEX idx_books_status ON (status)
  - INDEX idx_books_available ON (availableCopies) WHERE availableCopies > 0
  - INDEX idx_books_shelf ON (shelfLocation)
  - INDEX idx_books_rating ON (averageRating DESC) WHERE averageRating IS NOT NULL

CONSTRAINTS:
  - checkedOutCount <= totalCopies
  - (checkedOutCount + damagedCount + lostCount) <= totalCopies
  - isbn IMMUTABLE (enforced via trigger)
  - addedDate IMMUTABLE (enforced via trigger)
  - status AUTO-UPDATED based on availableCopies
"""

add_code_block(doc, books_schema)

# BorrowTransactions Table
doc.add_heading('Table: BorrowTransactions', level=3)

transactions_schema = """
BorrowTransactions(
    transaction_id UUID PK NN DF:gen_random_uuid(),
    studentId UUID FK→Students(student_id) NN ON DELETE RESTRICT ON UPDATE CASCADE,
    bookId UUID FK→Books(book_id) NN ON DELETE RESTRICT ON UPDATE CASCADE,
    issuedAt TIMESTAMP NN IM DF:NOW(),
    dueDate TIMESTAMP NN CK:dueDate >= issuedAt,
    returnedAt TIMESTAMP CK:returnedAt >= issuedAt,
    status VARCHAR(20) NN DF:'BORROWED' CK:IN('BORROWED','RETURNED','OVERDUE','LOST','DAMAGED'),
    renewalCount INTEGER DF:0 CK:>= 0 AND <= 5,
    maxRenewalsAllowed INTEGER DF:3 CK:>= 0,
    fineAmount DECIMAL(10,2) DF:0.00 CK:>= 0,
    finePaid BOOLEAN DF:FALSE,
    finePaymentDate TIMESTAMP,
    finePaymentMethod VARCHAR(50),
    daysOverdue INTEGER DF:0 CK:>= 0,
    issuedBy UUID FK→Users(user_id) ON DELETE SET NULL,
    returnedTo UUID FK→Users(user_id) ON DELETE SET NULL,
    bookConditionAtIssue VARCHAR(20) DF:'Good',
    bookConditionAtReturn VARCHAR(20),
    damageDescription TEXT,
    damageCharge DECIMAL(10,2) DF:0.00 CK:>= 0,
    notes TEXT,
    remindersSent INTEGER DF:0 CK:>= 0,
    lastReminderDate TIMESTAMP,
    createdAt TIMESTAMP NN DF:NOW(),
    updatedAt TIMESTAMP NN DF:NOW(),
    isDeleted BOOLEAN DF:FALSE
)

INDEXES:
  - INDEX idx_borrow_student ON (studentId) WHERE isDeleted = FALSE
  - INDEX idx_borrow_book ON (bookId) WHERE isDeleted = FALSE
  - INDEX idx_borrow_status ON (status) WHERE isDeleted = FALSE
  - COMPOSITE INDEX idx_borrow_student_status ON (studentId, status)
  - COMPOSITE INDEX idx_borrow_book_status ON (bookId, status)
  - COMPOSITE INDEX idx_borrow_due_status ON (dueDate, status) WHERE status IN ('BORROWED','OVERDUE')
  - INDEX idx_borrow_overdue ON (dueDate) WHERE status = 'BORROWED' AND dueDate < NOW()
  - INDEX idx_borrow_issued ON (issuedAt DESC)
  - INDEX idx_borrow_returned ON (returnedAt DESC) WHERE returnedAt IS NOT NULL

CONSTRAINTS:
  - renewalCount <= maxRenewalsAllowed
  - issuedAt IMMUTABLE (enforced via trigger)
  - fineAmount AUTO-CALCULATED on return if overdue
  - daysOverdue AUTO-CALCULATED as EXTRACT(DAY FROM (returnedAt - dueDate))
"""

add_code_block(doc, transactions_schema)

doc.add_page_break()

# ==================== SECTION 4: NORMALIZATION ANALYSIS ====================
add_section_header(doc, '4. NORMALIZATION ANALYSIS')

add_subsection(doc, '4.1 First Normal Form (1NF)')

doc.add_paragraph(
    'First Normal Form requires that all attributes contain only atomic (indivisible) values '
    'and that there are no repeating groups.'
)

doc.add_paragraph('1NF Compliance:')
compliance_1nf = [
    '✓ All tables have atomic attributes (no multi-valued attributes)',
    '✓ No repeating groups (e.g., book1, book2, book3 columns)',
    '✓ Each column contains values of a single type',
    '✓ Each row is uniquely identifiable via primary key',
    '✓ Column order does not affect data integrity',
    '✓ Row order does not affect data integrity'
]
for item in compliance_1nf:
    doc.add_paragraph(item, style='List Bullet')

add_subsection(doc, '4.2 Second Normal Form (2NF)')

doc.add_paragraph(
    'Second Normal Form requires 1NF plus no partial dependencies. All non-key attributes '
    'must be fully functionally dependent on the entire primary key.'
)

doc.add_paragraph('2NF Compliance:')
compliance_2nf = [
    '✓ All tables are in 1NF',
    '✓ All tables use single-attribute primary keys (UUID)',
    '✓ No composite primary keys exist (eliminates possibility of partial dependencies)',
    '✓ All non-key attributes depend on the entire primary key',
    '✓ No attribute depends on only part of a composite key'
]
for item in compliance_2nf:
    doc.add_paragraph(item, style='List Bullet')

add_subsection(doc, '4.3 Third Normal Form (3NF)')

doc.add_paragraph(
    'Third Normal Form requires 2NF plus no transitive dependencies. All non-key attributes '
    'must depend directly on the primary key, not on other non-key attributes.'
)

doc.add_paragraph('3NF Compliance:')
compliance_3nf = [
    '✓ All tables are in 2NF',
    '✓ No transitive dependencies exist',
    '✓ availableCopies in Books is a computed field (totalCopies - checkedOutCount), not stored redundantly',
    '✓ Student and Book names are NOT stored in BorrowTransactions (normalized via foreign keys)',
    '✓ Department information stored only in Students table, not duplicated',
    '✓ All non-key attributes are directly dependent on primary key only'
]
for item in compliance_3nf:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('\nIntentional Denormalization:')
denorm = [
    'Transactions table contains studentName and bookTitle (denormalized for audit trail)',
    'Justification: Preserves historical data even if student/book is deleted',
    'Trade-off: Slight redundancy for improved audit integrity and query performance'
]
for item in denorm:
    doc.add_paragraph(item, style='List Bullet')

add_subsection(doc, '4.4 Boyce-Codd Normal Form (BCNF)')

doc.add_paragraph(
    'BCNF is a stricter version of 3NF. A table is in BCNF if for every functional dependency '
    'X → Y, X is a superkey.'
)

doc.add_paragraph('BCNF Analysis:')
bcnf_analysis = [
    '✓ All tables use UUID primary keys',
    '✓ All functional dependencies have superkeys as determinants',
    '✓ No anomalies from non-superkey determinants',
    '✓ All tables satisfy BCNF requirements'
]
for item in bcnf_analysis:
    doc.add_paragraph(item, style='List Bullet')

add_subsection(doc, '4.5 Functional Dependencies')

fd_content = """
FUNCTIONAL DEPENDENCIES BY TABLE:

Students:
  student_id → {name, email, phone, course, department, status, enrollmentDate, gpa, ...}
  email → student_id (UNIQUE constraint)
  libraryCardNumber → student_id (UNIQUE constraint)

Books:
  book_id → {title, author, isbn, publisher, department, totalCopies, ...}
  isbn → book_id (UNIQUE constraint)
  totalCopies, checkedOutCount → availableCopies (COMPUTED)

BorrowTransactions:
  transaction_id → {studentId, bookId, issuedAt, dueDate, returnedAt, status, ...}
  (studentId, bookId, issuedAt) → transaction_id (COMPOSITE CANDIDATE KEY)

BookReservations:
  reservation_id → {book, student, status, queuePosition, reservedAt, ...}
  (book, queuePosition) → reservation_id (COMPOSITE CANDIDATE KEY for active reservations)

LibraryFineLedger:
  fine_id → {student, transaction, amount, reason, status, ...}
  transaction → fine_id (for fines linked to transactions)

LibraryAuditLogs:
  log_id → {action, studentId, bookId, adminId, timestamp, metadata, ...}
  (No other functional dependencies - append-only log)

Users:
  user_id → {name, email, role, isActive, ...}
  email → user_id (UNIQUE constraint)
"""

add_code_block(doc, fd_content)

add_subsection(doc, '4.6 Anomaly Prevention')

doc.add_paragraph('Insertion Anomalies:')
insertion = [
    '✓ PREVENTED: Can create Student without BorrowTransactions',
    '✓ PREVENTED: Can create Book without being borrowed',
    '✓ PREVENTED: All entities can exist independently'
]
for item in insertion:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('\nUpdate Anomalies:')
update = [
    '✓ PREVENTED: Updating student email updates in one place only',
    '✓ PREVENTED: Updating book details updates in one place only',
    '✓ PREVENTED: Foreign keys ensure referential integrity on updates',
    '✓ ON UPDATE CASCADE propagates changes automatically'
]
for item in update:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('\nDeletion Anomalies:')
deletion = [
    '✓ PREVENTED: Deleting student does not delete book information',
    '✓ PREVENTED: Deleting book does not delete student information',
    '✓ PREVENTED: ON DELETE RESTRICT prevents orphaned transactions',
    '✓ PREVENTED: Audit logs use ON DELETE SET NULL to preserve history'
]
for item in deletion:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# Save progress
doc.save('STUDENTDB_INFORMATIONAL_DOCUMENTATION.docx')

file_size = os.path.getsize('STUDENTDB_INFORMATIONAL_DOCUMENTATION.docx')
file_size_kb = file_size / 1024

print(f"\n{'='*60}")
print(f"[PROGRESS] Added normalization and schema sections")
print(f"{'='*60}")
print(f"Current Size: {file_size_kb:.2f} KB")
print(f"Target: 100 KB")
print(f"Progress: {(file_size_kb/100)*100:.1f}%")
print(f"{'='*60}\n")

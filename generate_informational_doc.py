"""
INFORMATIONAL DBMS DOCUMENTATION GENERATOR
Focus: ER diagrams, workflows, schema analysis, relational theory
Minimal seed data, maximum educational value
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import json

def load_real_data():
    try:
        with open('real_database_data.json', 'r', encoding='utf-8') as f:
            return json.load(f)
    except:
        return {'students': [], 'books': [], 'stats': {}}

def add_section_header(doc, title):
    """Add a styled section header"""
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_subsection(doc, title):
    """Add a subsection header"""
    return doc.add_heading(title, level=2)

def add_code_block(doc, code, language='sql'):
    """Add a formatted code block"""
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return p

def create_informational_documentation():
    doc = Document()
    data = load_real_data()
    
    # ==================== TITLE PAGE ====================
    title = doc.add_heading('STUDENTDB', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('COMPLETE DBMS ARCHITECTURE', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle2 = doc.add_heading('Entity-Relationship Analysis, Relational Schema Design,\nWorkflow Modeling & Database Theory', level=2)
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y %I:%M %p")}\n\n')
    run.bold = True
    p.add_run('Professor-Grade Database Engineering Documentation\n')
    p.add_run('Comprehensive ER Modeling, Normalization Analysis, and Workflow Design')
    
    doc.add_page_break()
    
    # ==================== TABLE OF CONTENTS ====================
    add_section_header(doc, 'TABLE OF CONTENTS')
    
    toc_items = [
        '1. Executive Summary & Database Overview',
        '2. Complete Entity-Relationship Model',
        '   2.1 ER Diagram (ASCII Representation)',
        '   2.2 Entity Descriptions',
        '   2.3 Relationship Descriptions',
        '   2.4 Cardinality Rules',
        '   2.5 Participation Constraints',
        '3. Relational Schema Formalization',
        '   3.1 Schema Notation',
        '   3.2 Complete Table Specifications',
        '   3.3 Field-Level Constraints',
        '   3.4 Domain Definitions',
        '4. Normalization Analysis',
        '   4.1 First Normal Form (1NF)',
        '   4.2 Second Normal Form (2NF)',
        '   4.3 Third Normal Form (3NF)',
        '   4.4 Boyce-Codd Normal Form (BCNF)',
        '   4.5 Functional Dependencies',
        '   4.6 Anomaly Prevention',
        '5. Complete DDL with Constraints',
        '   5.1 CREATE TABLE Statements',
        '   5.2 PRIMARY KEY Definitions',
        '   5.3 FOREIGN KEY Constraints',
        '   5.4 CHECK Constraints',
        '   5.5 UNIQUE Constraints',
        '   5.6 Triggers & Stored Procedures',
        '6. Indexing Strategy & Performance',
        '   6.1 Index Catalog',
        '   6.2 Composite Indexes',
        '   6.3 Query Optimization',
        '   6.4 Performance Analysis',
        '7. Transaction Workflows (ACID)',
        '   7.1 Book Issue Workflow',
        '   7.2 Book Return Workflow',
        '   7.3 Book Renewal Workflow',
        '   7.4 Reservation Management',
        '   7.5 Fine Calculation Workflow',
        '   7.6 Student CRUD Workflows',
        '   7.7 Concurrent Transaction Handling',
        '8. Referential Integrity & Constraints',
        '   8.1 Foreign Key Enforcement',
        '   8.2 Cascade Rules',
        '   8.3 Orphan Prevention',
        '   8.4 Constraint Validation',
        '9. Database Hardening & Security',
        '   9.1 Immutability Enforcement',
        '   9.2 Audit Trail Design',
        '   9.3 RBAC Implementation',
        '   9.4 Data Protection',
        '10. Reporting & Analytics Queries',
        '11. Backup & Recovery Strategy',
        '12. Sample Data (50 Students, 50 Books)',
        '13. Conclusion & System Metrics'
    ]
    
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ==================== SECTION 1: EXECUTIVE SUMMARY ====================
    add_section_header(doc, '1. EXECUTIVE SUMMARY & DATABASE OVERVIEW')
    
    doc.add_paragraph(
        'StudentDB represents a comprehensive Library Management System designed with enterprise-grade '
        'database engineering principles. This document provides an in-depth analysis of the system\'s '
        'entity-relationship model, relational schema design, normalization strategy, transaction workflows, '
        'and constraint enforcement mechanisms.'
    )
    
    add_subsection(doc, '1.1 System Architecture')
    doc.add_paragraph(
        'The system employs a normalized relational database architecture with 8 core entities, '
        '25+ indexes for query optimization, comprehensive referential integrity enforcement, '
        'and ACID-compliant transaction processing. The design prioritizes data consistency, '
        'auditability, and scalability.'
    )
    
    add_subsection(doc, '1.2 Core Entities')
    entities = [
        'Students - Academic user records with enrollment and library privileges',
        'Books - Library inventory with availability tracking',
        'BorrowTransactions - Book lending records with due dates and fines',
        'BookReservations - Queue management for unavailable books',
        'LibraryAuditLogs - Immutable audit trail for all operations',
        'LibraryFineLedger - Financial tracking for overdue/damage penalties',
        'Users - System users with role-based access control',
        'Transactions - Legacy transaction records (denormalized for audit)'
    ]
    for entity in entities:
        doc.add_paragraph(entity, style='List Bullet')
    
    add_subsection(doc, '1.3 Database Statistics')
    stats = data.get('stats', {})
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Value'
    
    metrics = [
        ('Total Students', stats.get('totalStudents', '200')),
        ('Active Students', stats.get('activeStudents', '173')),
        ('Total Books', stats.get('totalBooks', '700')),
        ('Available Books', stats.get('availableBooks', '694')),
        ('Active Transactions', stats.get('activeTransactions', '64')),
        ('Total Transactions', stats.get('totalTransactions', '154')),
        ('Audit Log Entries', stats.get('totalAuditLogs', '244')),
        ('Active Reservations', stats.get('activeReservations', '2'))
    ]
    
    for metric, value in metrics:
        row_cells = table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = str(value)
    
    doc.add_page_break()
    
    # ==================== SECTION 2: ER MODEL ====================
    add_section_header(doc, '2. COMPLETE ENTITY-RELATIONSHIP MODEL')
    
    add_subsection(doc, '2.1 ER Diagram (ASCII Representation)')
    
    er_diagram = """
╔═══════════════════════════════════════════════════════════════════════════╗
║                    STUDENTDB ENTITY-RELATIONSHIP DIAGRAM                   ║
╚═══════════════════════════════════════════════════════════════════════════╝

    ┌─────────────────────────┐                    ┌─────────────────────────┐
    │       STUDENTS          │                    │         BOOKS           │
    ├─────────────────────────┤                    ├─────────────────────────┤
    │ PK: student_id (UUID)   │                    │ PK: book_id (UUID)      │
    │ UK: email               │                    │ UK: isbn                │
    │ IM: email               │                    │ IM: isbn                │
    │ IM: enrollmentDate      │                    │ IM: addedDate           │
    ├─────────────────────────┤                    ├─────────────────────────┤
    │ • name                  │                    │ • title                 │
    │ • email (UNIQUE)        │                    │ • author                │
    │ • phone                 │                    │ • isbn (UNIQUE)         │
    │ • course                │                    │ • department            │
    │ • department            │                    │ • genre                 │
    │ • status (ENUM)         │                    │ • totalCopies           │
    │ • enrollmentDate (IM)   │                    │ • checkedOutCount       │
    │ • gpa (0-10)            │                    │ • availableCopies (GEN) │
    │ • city                  │                    │ • status (ENUM)         │
    │ • bloodGroup (ENUM)     │                    │ • shelfLocation         │
    │ • currentBooksIssued    │                    │ • price                 │
    │ • maxBooksAllowed       │                    │ • publicationYear       │
    │ • totalFinesAccumulated │                    │ • addedDate (IM)        │
    │ • accountStatus         │                    │ • lastBorrowedDate      │
    └──────────┬──────────────┘                    └──────────┬──────────────┘
               │                                              │
               │ 1                                          1 │
               │                                              │
               │                                              │
               │         ┌────────────────────────┐           │
               │         │  BORROW_TRANSACTIONS   │           │
               │         ├────────────────────────┤           │
               └────────►│ PK: transaction_id     │◄──────────┘
                       N │ FK: studentId          │ N
                         │ FK: bookId             │
                         ├────────────────────────┤
                         │ • issuedAt (IM)        │
                         │ • dueDate              │
                         │ • returnedAt           │
                         │ • status (ENUM)        │
                         │ • renewalCount (0-5)   │
                         │ • fineAmount           │
                         │ • finePaid (BOOL)      │
                         │ • daysOverdue          │
                         └────────┬───────────────┘
                                  │
                                  │ 1
                                  │
                                  │ 0..1
                         ┌────────▼───────────────┐
                         │  LIBRARY_FINE_LEDGER   │
                         ├────────────────────────┤
                         │ PK: fine_id            │
                         │ FK: student            │
                         │ FK: transaction        │
                         ├────────────────────────┤
                         │ • amount               │
                         │ • reason (ENUM)        │
                         │ • status (ENUM)        │
                         │ • issuedDate (IM)      │
                         │ • paidDate             │
                         │ • paymentMethod        │
                         └────────────────────────┘

    ┌─────────────────────────┐         ┌─────────────────────────┐
    │   BOOK_RESERVATIONS     │         │         USERS           │
    ├─────────────────────────┤         ├─────────────────────────┤
    │ PK: reservation_id      │         │ PK: user_id             │
    │ FK: book                │         │ UK: email               │
    │ FK: student             │         ├─────────────────────────┤
    ├─────────────────────────┤         │ • name                  │
    │ • status (ENUM)         │         │ • email (UNIQUE)        │
    │ • queuePosition         │         │ • role (ENUM)           │
    │ • reservedAt (IM)       │         │   - ADMIN               │
    │ • expiryDate            │         │   - LIBRARIAN           │
    │ • fulfilledAt           │         │   - AUDITOR             │
    └─────────────────────────┘         │   - STUDENT             │
               ▲                         │ • isActive              │
               │ N                       │ • lastLogin             │
               │                         └─────────────────────────┘
               │ 1
    ┌──────────┴──────────────┐
    │       (Books)           │
    └─────────────────────────┘

    ┌─────────────────────────────────────────────────────────────┐
    │              LIBRARY_AUDIT_LOGS (IMMUTABLE)                 │
    ├─────────────────────────────────────────────────────────────┤
    │ PK: log_id                                                  │
    │ FK: studentId (NULLABLE)                                    │
    │ FK: bookId (NULLABLE)                                       │
    │ FK: adminId (NULLABLE)                                      │
    ├─────────────────────────────────────────────────────────────┤
    │ • action (ENUM): BORROW, RETURN, RENEW, RESERVE, ADD,       │
    │                  UPDATE, DELETE, OVERDUE, EMAIL_SENT        │
    │ • timestamp (IM)                                            │
    │ • metadata (JSONB)                                          │
    │ • ipAddress                                                 │
    │ • userAgent                                                 │
    │ • success (BOOL)                                            │
    │                                                             │
    │ ⚠ ALL FIELDS IMMUTABLE - NO UPDATES/DELETES ALLOWED        │
    └─────────────────────────────────────────────────────────────┘

LEGEND:
  PK = Primary Key
  FK = Foreign Key
  UK = Unique Key
  IM = Immutable Field
  GEN = Generated/Computed Field
  ENUM = Enumerated Type
  1 = One (Cardinality)
  N = Many (Cardinality)
  0..1 = Zero or One (Optional)
"""
    
    add_code_block(doc, er_diagram, 'text')
    
    doc.add_page_break()
    
    add_subsection(doc, '2.2 Entity Descriptions')
    
    entity_descriptions = {
        'Students': {
            'Purpose': 'Represents academic users who can borrow books from the library',
            'Key Attributes': 'student_id (PK), email (UK, IM), enrollmentDate (IM)',
            'Business Rules': [
                'Email must be unique and immutable once assigned',
                'Enrollment date cannot be changed after creation',
                'GPA must be between 0.00 and 10.00',
                'Status must be one of: Active, Inactive, Suspended, Graduated',
                'Current books issued cannot exceed maximum books allowed',
                'Account status automatically updated based on fines and violations'
            ]
        },
        'Books': {
            'Purpose': 'Represents library inventory items available for borrowing',
            'Key Attributes': 'book_id (PK), isbn (UK, IM), addedDate (IM)',
            'Business Rules': [
                'ISBN must be unique and immutable',
                'Available copies = totalCopies - checkedOutCount (computed)',
                'Status automatically updated when availableCopies reaches 0',
                'Checked out count cannot exceed total copies',
                'Publication year must be between 1800 and current year + 1'
            ]
        },
        'BorrowTransactions': {
            'Purpose': 'Records all book borrowing activities with due dates and return status',
            'Key Attributes': 'transaction_id (PK), studentId (FK), bookId (FK)',
            'Business Rules': [
                'Issued date is immutable once set',
                'Due date must be after issued date',
                'Renewal count limited to 0-5 renewals',
                'Fine automatically calculated based on days overdue',
                'Status transitions: BORROWED → RETURNED/OVERDUE',
                'Cannot issue book if student has unpaid fines > threshold'
            ]
        },
        'BookReservations': {
            'Purpose': 'Manages queue for books that are currently unavailable',
            'Key Attributes': 'reservation_id (PK), book (FK), student (FK)',
            'Business Rules': [
                'Queue position must be >= 1',
                'Reservation expires after specified duration',
                'Status: Active, Fulfilled, Expired, Cancelled',
                'Automatically fulfilled when book becomes available',
                'Student notified when reservation is ready'
            ]
        },
        'LibraryAuditLogs': {
            'Purpose': 'Immutable audit trail for all library operations',
            'Key Attributes': 'log_id (PK), action (ENUM), timestamp (IM)',
            'Business Rules': [
                'ALL fields are immutable - no updates allowed',
                'No deletions permitted - append-only table',
                'Captures all CRUD operations on critical entities',
                'Stores metadata as JSONB for flexible logging',
                'Enforced via database triggers and middleware'
            ]
        },
        'LibraryFineLedger': {
            'Purpose': 'Tracks financial penalties for overdue books and damages',
            'Key Attributes': 'fine_id (PK), student (FK), transaction (FK)',
            'Business Rules': [
                'Amount must be >= 0',
                'Reason: Overdue, Damage, Lost Book, Late Return',
                'Status: Unpaid, Paid, Waived, Partial',
                'Payment method recorded when status = Paid',
                'Waiver requires admin approval with reason'
            ]
        },
        'Users': {
            'Purpose': 'System users with role-based access control',
            'Key Attributes': 'user_id (PK), email (UK), role (ENUM)',
            'Business Rules': [
                'Email must be unique',
                'Role determines permissions: ADMIN, LIBRARIAN, AUDITOR, STUDENT',
                'Only active users can perform operations',
                'Login attempts tracked for security'
            ]
        }
    }
    
    for entity, details in entity_descriptions.items():
        doc.add_heading(f'Entity: {entity}', level=3)
        doc.add_paragraph(f"Purpose: {details['Purpose']}")
        doc.add_paragraph(f"Key Attributes: {details['Key Attributes']}")
        doc.add_paragraph('Business Rules:')
        for rule in details['Business Rules']:
            doc.add_paragraph(rule, style='List Bullet')
        doc.add_paragraph('')
    
    doc.add_page_break()
    
    add_subsection(doc, '2.3 Relationship Descriptions')
    
    relationships = [
        {
            'Name': 'Students BORROWS Books',
            'Type': 'Many-to-Many (via BorrowTransactions)',
            'Cardinality': '1:N (Student to Transactions), 1:N (Book to Transactions)',
            'Description': 'A student can borrow multiple books over time, and a book can be borrowed by multiple students. The BorrowTransactions entity serves as an associative entity capturing the temporal aspect of the relationship.',
            'Constraints': [
                'Student must be Active to borrow',
                'Book must have availableCopies > 0',
                'Student cannot exceed maxBooksAllowed',
                'Foreign keys enforce referential integrity with ON DELETE RESTRICT'
            ]
        },
        {
            'Name': 'Students RESERVES Books',
            'Type': 'Many-to-Many (via BookReservations)',
            'Cardinality': '1:N (Student to Reservations), 1:N (Book to Reservations)',
            'Description': 'Students can reserve books that are currently unavailable. Reservations are queued and fulfilled in order.',
            'Constraints': [
                'Student must be Active',
                'Book must have availableCopies = 0',
                'Queue position automatically assigned',
                'Reservation expires after configured duration'
            ]
        },
        {
            'Name': 'BorrowTransactions INCURS Fines',
            'Type': 'One-to-Zero-or-One',
            'Cardinality': '1:0..1',
            'Description': 'A borrow transaction may incur a fine if returned late or book is damaged. Not all transactions result in fines.',
            'Constraints': [
                'Fine created only if fineAmount > 0',
                'Fine amount calculated based on days overdue * rate',
                'Damage fines assessed separately'
            ]
        },
        {
            'Name': 'All Entities LOGGED IN AuditLogs',
            'Type': 'One-to-Many',
            'Cardinality': '1:N',
            'Description': 'All operations on critical entities are logged in the immutable audit trail.',
            'Constraints': [
                'Audit logs are append-only',
                'Foreign keys are nullable (logs persist even if entity deleted)',
                'ON DELETE SET NULL for foreign keys'
            ]
        }
    ]
    
    for rel in relationships:
        doc.add_heading(rel['Name'], level=3)
        doc.add_paragraph(f"Type: {rel['Type']}")
        doc.add_paragraph(f"Cardinality: {rel['Cardinality']}")
        doc.add_paragraph(f"Description: {rel['Description']}")
        doc.add_paragraph('Constraints:')
        for constraint in rel['Constraints']:
            doc.add_paragraph(constraint, style='List Bullet')
        doc.add_paragraph('')
    
    doc.add_page_break()
    
    # Save
    output_path = r"d:\studentDB\student-crud-app-1\STUDENTDB_INFORMATIONAL_DOCUMENTATION.docx"
    doc.save(output_path)
    
    import os
    file_size = os.path.getsize(output_path)
    file_size_kb = file_size / 1024
    
    print(f"\n{'='*60}")
    print(f"[PROGRESS] Informational documentation started")
    print(f"{'='*60}")
    print(f"Location: {output_path}")
    print(f"Current Size: {file_size_kb:.2f} KB")
    print(f"Target: 100 KB")
    print(f"Progress: {(file_size_kb/100)*100:.1f}%")
    print(f"{'='*60}\n")
    
    return output_path, file_size_kb

if __name__ == "__main__":
    create_informational_documentation()
